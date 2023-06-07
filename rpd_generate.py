from app import app
from docxtpl import DocxTemplate
import io
from docx.shared import Pt
from string import Template
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from flask import render_template, request, Response, jsonify
from pymongo import MongoClient
from bson import ObjectId
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml.shared import qn
from bson import ObjectId



client = MongoClient("mongodb+srv://4enpavel:NdEKoRqD6xlQGHUB@planeducation.jaybzzr.mongodb.net/test?retryWrites=true&w=majority")
db = client.test

@app.route('/rpd_generate')
def rpd_generate():
    # Получение списка уникальных годов из базы данных
    years = db.disciplines.distinct('year_education')
    available_rooms = db.rooms.distinct('room')
    years.insert(0, "Выберети год")
    return render_template('rpd_generate.html', years=years, available_rooms=available_rooms)

@app.route('/get_educational_programs', methods=['POST'])
def get_educational_programs():
    year = request.form['year']
    programs = db.disciplines.distinct("name_education", {"year_education": year})
    programs.insert(0, "Выберите образовательную программу")
    return {'programs': programs}

@app.route('/get_directions', methods=['POST'])
def get_directions():
    program = request.form['program']
    directions = db.disciplines.distinct("name_direction", {"name_education": program})
    directions.insert(0, "Выберите направление подготовки")
    return {'directions': directions}

@app.route('/get_disciplines', methods=['POST'])
def get_disciplines():
    direction = request.form['direction']
    year = request.form['year']
    disciplines = db.disciplines.distinct("discipline_name", {"name_direction": direction, "year_education":year})
    disciplines.insert(0, "Выберети дисциплину")


    return {'disciplines': disciplines}


@app.route('/get_department', methods=['POST'])
def get_department():
    discipline = request.form['discipline']
    name_direction = request.form['direction']  # Получаем выбранное направление подготовки
    department = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['department_name']
    return {'department': department}

@app.route('/get_school', methods=['POST'])
def get_school():
    discipline = request.form['discipline']
    name_direction = request.form['direction']  # Получаем выбранное направление подготовки
    school = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['school_name']
    return {'school': school}

@app.route('/get_form_education', methods=['POST'])
def get_form_education():
    discipline = request.form['discipline']
    name_direction = request.form['direction']  # Получаем выбранное направление подготовки
    form_education = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['form_educ']
    return {'form_education': form_education}

@app.route('/get_course', methods=['POST'])
def get_course():
    discipline = request.form['discipline']
    name_direction = request.form['direction']  # Получаем выбранное направление подготовки
    course = db.disciplines.distinct("work_hours.course",{"discipline_name": discipline, "name_direction": name_direction})
    course = ', '.join(map(str, course))
    return {'course': course}

@app.route('/get_semester', methods=['POST'])
def get_semester():
    discipline = request.form['discipline']
    name_direction = request.form['direction']  # Получаем выбранное направление подготовки
    semester = db.disciplines.distinct("work_hours.semestr",{"discipline_name": discipline, "name_direction": name_direction})
    semester = ', '.join(map(str, semester))
    return {'semester': semester}

@app.route('/get_titul_name', methods=['POST'])
def get_titul_name():
    discipline = request.form['discipline']
    name_direction = request.form['direction']  # Получаем выбранное направление подготовки
    titul_name = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['titul_name']
    return {'titul_name': titul_name}


@app.route('/get_date_aprooval', methods=['POST'])
def get_date_aprooval():
    discipline = request.form['discipline']
    name_direction = request.form['direction']  # Получаем выбранное направление подготовки
    date_aprooval = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['date_approoval']
    return {'date_aprooval': date_aprooval}

@app.route('/get_number_fgos', methods=['POST'])
def get_number_fgos():
    discipline = request.form['discipline']
    name_direction = request.form['direction']  # Получаем выбранное направление подготовки
    number_fgos = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['number_fgos']
    return {'number_fgos': number_fgos}

@app.route('/get_code_education', methods=['POST'])
def get_code_education():
    discipline = request.form['discipline']
    name_direction = request.form['direction']  # Получаем выбранное направление подготовки
    code_education = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['code_education']
    return {'code_education': code_education}

@app.route('/get_lecture', methods=['POST'])
def get_lecture():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    work_hours = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    lecture_hours = 0
    if work_hours:
        for item in work_hours.get("work_hours", []):
            if item.get("work_type") == "Лекционные занятия":
                lecture_hours += int(item.get("hours", 0))
    
    return {'lecture': lecture_hours}

@app.route('/get_laborator', methods=['POST'])
def get_laborator():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    work_hours = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    laborator_hours = 0
    if work_hours:
        for item in work_hours.get("work_hours", []):
            if item.get("work_type") == "Лабораторные занятия":
                laborator_hours += int(item.get("hours", 0))
    
    return {'laborator': laborator_hours}

@app.route('/get_zachet_edenic', methods=['POST'])
def get_zachet_edenic():
    discipline = request.form['discipline']
    name_direction = request.form['direction']  # Получаем выбранное направление подготовки
    zachet_edenic = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['zachet_edenic']
    return {'zachet_edenic': zachet_edenic}


@app.route('/get_practice', methods=['POST'])
def get_practice():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    work_hours = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    practice_hours = 0
    if work_hours:
        for item in work_hours.get("work_hours", []):
            if item.get("work_type") == "Практические занятия":
                practice_hours += int(item.get("hours", 0))
    
    return {'practice': practice_hours}

@app.route('/get_sam_work', methods=['POST'])
def get_sam_work():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    work_hours = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    sam_work_hours = 0
    if work_hours:
        for item in work_hours.get("work_hours", []):
            if item.get("work_type") == "Самостоятельная работа":
                sam_work_hours += int(item.get("hours", 0))
    
    return {'sam_work': sam_work_hours}

@app.route('/get_control', methods=['POST'])
def get_control():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    work_hours = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    control_hours = 0
    if work_hours:
        for item in work_hours.get("work_hours", []):
            if item.get("work_type") == "Часы на контроль":
                control_hours += int(item.get("hours", 0))
    
    return {'control': control_hours}

@app.route('/get_competencies', methods=['POST'])
def get_competencies():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    result = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    competency_names = []
    if result:
        competencies = result.get("competencies", [])
        for competency in competencies:
            competency_name = competency.get("competency_name")
            if competency_name:
                competency_names.append(competency_name)
    
    competency_names_str = ', '.join(competency_names)
    
    return competency_names_str

@app.route('/get_competencies_table', methods=['POST'])
def get_competencies_table():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    result = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    competency_data = []
    if result:
        competencies = result.get("competencies", [])
        for index, competency in enumerate(competencies):
            new_competency_name = competency.get("new_competency_name")
            new_competency_description = competency.get("new_competency_description")
            new_competency_info = new_competency_name + " " + new_competency_description
            competency_name = competency.get("competency_name")
            competency_description = competency.get("competency_description")
            competency_info = competency_name + " " + competency_description
            if new_competency_name and competency_name:
                competency_type = competency.get("type_competitions")
                competency_data.append({
                    # 'index': index + 1,
                    'new_competency_name': new_competency_info,
                    'competency_name': competency_info,
                    'type_competitions': competency_type
                })
    
    return jsonify(competency_data)

@app.route('/get_uk_competencies', methods=['POST'])
def get_uk_competencies():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    result = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    competency_data = []
    if result:
        competencies = result.get("competencies", [])
        for index, competency in enumerate(competencies):
            new_competency_name = competency.get("new_competency_name")
            if new_competency_name and new_competency_name.startswith("УК-"):
                new_competency_description = competency.get("new_competency_description")
                new_competency_info = new_competency_name + " " + new_competency_description
                competency_name = competency.get("competency_name")
                competency_description = competency.get("competency_description")
                competency_info = competency_name + " " + competency_description
                competency_type = competency.get("type_competitions")
                competency_data.append({
                    'new_competency_name': new_competency_info,
                    'competency_name': competency_info,
                    'type_competitions': competency_type
                })
    
    return jsonify(competency_data)

@app.route('/get_opk_competencies', methods=['POST'])
def get_opk_competencies():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    result = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    competency_data = []
    if result:
        competencies = result.get("competencies", [])
        for index, competency in enumerate(competencies):
            new_competency_name = competency.get("new_competency_name")
            if new_competency_name and new_competency_name.startswith("ОПК-"):
                new_competency_description = competency.get("new_competency_description")
                new_competency_info = new_competency_name + " " + new_competency_description
                competency_name = competency.get("competency_name")
                competency_description = competency.get("competency_description")
                competency_info = competency_name + " " + competency_description
                competency_type = competency.get("type_competitions")
                competency_data.append({
                    'new_competency_name': new_competency_info,
                    'competency_name': competency_info,
                    'type_competitions': competency_type
                })
    
    return jsonify(competency_data)

@app.route('/get_pk_competencies', methods=['POST'])
def get_pk_competencies():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    result = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    competency_data = []
    if result:
        competencies = result.get("competencies", [])
        for index, competency in enumerate(competencies):
            new_competency_name = competency.get("new_competency_name")
            if new_competency_name and new_competency_name.startswith("ПК-"):
                new_competency_description = competency.get("new_competency_description")
                new_competency_info = new_competency_name + " " + new_competency_description
                competency_name = competency.get("competency_name")
                competency_description = competency.get("competency_description")
                competency_info = competency_name + " " + competency_description
                competency_type = competency.get("type_competitions")
                competency_data.append({
                    'new_competency_name': new_competency_info,
                    'competency_name': competency_info,
                    'type_competitions': competency_type
                })
    
    return jsonify(competency_data)

@app.route('/filter', methods=['POST'])
def filter_rooms():
    selected_rooms = request.form.getlist('room')

    # Проверка, есть ли выбранные комнаты
    if len(selected_rooms) == 0:
        return jsonify([])

    # Запрос данных комнат из базы данных по выбранным комнатам
    rooms = db.rooms.find({'room': {'$in': selected_rooms}})

    # Формирование списка данных комнат
    room_list = []
    for room in rooms:
        room_data = {
            'room_description': room['room_description'],
            'room_seets': room['room_seets'],
            'room_device': room['room_device']
        }
        room_list.append(room_data)

    return jsonify(room_list)
    

@app.route('/get_form_control', methods=['POST'])
def get_form_control():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    work_hours = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    form_control = set()

    if work_hours:
        for item in work_hours.get("work_hours", []):
            work_type = item.get("work_type")
            if work_type in ['Экзамен', 'Эссе', 'Зачет', 'Зачет с оценкой', 'Курсовой проект', 'Курсовая работа',
                             'Контрольная работа', 'Домашняя контрольная работа', 'Оценка', 'Реферат',
                             'Расчетно-графическая работа']:
                if work_type == 'Экзамен':
                    form_control.add('Экзаменом')
                elif work_type == 'Курсовая работа':
                    form_control.add('Курсовой работой')
                elif work_type == 'Зачет':
                    form_control.add('Зачётом')
                elif work_type == 'Расчетно-графическая работа':
                    form_control.add('Расчётно-графической работой')
                elif work_type == 'Зачет с оценкой':
                    form_control.add('Зачётом с оценкой')
                elif work_type == 'Курсовой проект':
                    form_control.add('Курсовым проектом')
                elif work_type == 'Контрольная работа':
                    form_control.add('Контрольной работой')
                elif work_type == 'Эссе':
                    form_control.add('Эссе')
                elif work_type == 'Реферат':
                    form_control.add('Рефератом')
                elif work_type == 'Оценка':
                    form_control.add('Оценкой')
                elif work_type == 'Домашняя контрольная работа':
                    form_control.add('Домашней контрольной работой')
    
    form_control_str = ', '.join(form_control)
    
    return form_control_str

@app.route('/get_form_control1', methods=['POST'])
def get_form_control1():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    
    work_hours = db.disciplines.find_one(
        {"discipline_name": discipline, "name_direction": name_direction}
    )

    form_control = set()

    if work_hours:
        for item in work_hours.get("work_hours", []):
            work_type = item.get("work_type")
            if work_type in ['Экзамен', 'Эссе', 'Зачет', 'Зачет с оценкой', 'Курсовой проект', 'Курсовая работа',
                             'Контрольная работа', 'Домашняя контрольная работа', 'Оценка', 'Реферат',
                             'Расчетно-графическая работа']:
                form_control.add(work_type)
    
    form_control_str = ', '.join(form_control)
    
    return jsonify(form_control=form_control_str)


@app.route('/get_room_types', methods=['GET'])
def get_room_types():
    room_types = db.rooms.distinct('room_type')
    return jsonify({'room_types': room_types})

@app.route('/get_rooms', methods=['GET'])
def get_rooms():
    room_type = request.args.get('room_type')
    rooms = db.rooms.find({'room_type': room_type})

    # Convert ObjectId to string
    serialized_rooms = []
    for room in rooms:
        room['_id'] = str(room['_id'])
        serialized_rooms.append(room)

    return jsonify({'rooms': serialized_rooms})



@app.route('/get_programms', methods=['GET'])
def get_programms():
    programms = db.programms.distinct('name_po')
    return jsonify({'programms': programms})

@app.route('/generate_document', methods=['POST'])
def generate_document():
    discipline = request.form['discipline']
    name_direction = request.form['direction']
    year = request.form['year']
    program = request.form['program']
    main_literature = request.form['main_literature']
    additional_literature = request.form['additional_literature']
    internet_literature = request.form['internet_literature']
    it_literature = request.form['it_literature']
    department = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['department_name']
    number_fgos = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['number_fgos']
    form_education = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['form_educ']
    date_aprooval = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['date_approoval']
    titul_name = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['titul_name']
    code_education = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['code_education']
    school = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['school_name']
    zachet_edenic = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['zachet_edenic']
    akadem_hour = db.disciplines.find_one({"discipline_name": discipline, "name_direction": name_direction})['akadem_hour']
    course = db.disciplines.distinct("work_hours.course",{"discipline_name": discipline, "name_direction": name_direction})
    course = ', '.join(map(str, course)).replace('[', '').replace(']', '').replace("'", "")
    language_discipline = request.form['language_discipline']
    form_control = get_form_control()
    lecture_hour = get_lecture()['lecture']
    practice_hour = get_practice()['practice']
    sm_hour = get_sam_work()['sam_work']
    authors = [request.form[key] for key in request.form if key.startswith('author')]
    authors = '\n'.join(authors)
    
    

    tasks_discipline = [request.form[key] for key in request.form if key.startswith('tasks-discipline')]
    formatted_tasks = ""
    for i, task in enumerate(tasks_discipline, start=1):
        formatted_tasks += f"{i})\xa0{task}\n"

    list_competencies = get_competencies()
    department_words = department.split(' ')
    first_word = department_words[0]
    target_discipline = request.form['target_discipline']

    theme_discipline = [request.form[key] for key in request.form if key.startswith('theme-discipline')]
    theme_discipline = " ".join(theme_discipline)

    theme_discipline_description = [request.form[key] for key in request.form if key.startswith('theme-discipline-description')]
    theme_discipline_description = ' '.join(theme_discipline_description)
    # formatted_text_theme = f"{theme_discipline}\n{theme_discipline_description}\n"

    if first_word == 'Департамент':
        first_word += 'а'
    elif first_word == 'Кафедра':
        first_word = first_word[:-1] + 'ы'
    elif first_word == 'Базовая' and department_words[1] == 'кафедра':
        first_word = 'Базовой'
        department_words[1] = 'кафедры'
    elif first_word == 'Академия':
        first_word = 'Академии'

    department_words[0] = first_word
    department = ' '.join(department_words)
    school_words = school.split(' ')
    first_word_school = school_words[0]
    if first_word_school == "Школа":
        first_word_school = "Школы"
    elif first_word_school == "Институт":
        first_word_school = "Института"

    school_words[0] = first_word_school
    school = ' '.join(school_words)

    # Загрузка шаблона документа Word
    document = DocxTemplate('RPD_Template.docx')
    # Получение данных таблицы из запроса
    column_data_fullcomp = []
    column_data_ukcomp = []
    column_data_opkcomp = []
    column_data_pkcomp = []
    column_data_hours = []
    column_data_comp = []
    column_data_mto = []


    column_names_fullcomp = ['column_comp1[]','column_comp2[]','column_comp3[]','column_comp4[]']
    column_names_ukcomp = ['column_comp_uk1[]','column_comp_uk2[]','column_comp_uk3[]','column_comp_uk4[]']
    column_names_opkcomp = ['column_comp_opk1[]','column_comp_opk2[]','column_comp_opk3[]','column_comp_opk4[]']
    column_names_pkcomp = ['column_comp_pk1[]','column_comp_pk2[]','column_comp_pk3[]','column_comp_pk4[]']
    column_names_hours = ['column1[]', 'column2[]', 'column3[]', 'column4[]', 'column5[]', 
                    'column6[]', 'column7[]', 'column8[]', 'column9[]', 'column10[]']
    column_names_comp = ['comp_col1[]','comp_col2[]','comp_col3[]','comp_col4[]','comp_col5[]']
    column_names_mto = ['mto_col1[]','mto_col2[]','mto_col3[]']
    

    for key in request.form.keys():
        if key in column_names_fullcomp:
            column_data_fullcomp.append(request.form.getlist(key))
            
    for key in request.form.keys():
        if key in column_names_ukcomp:
            column_data_ukcomp.append(request.form.getlist(key))

    for key in request.form.keys():
        if key in column_names_opkcomp:
            column_data_opkcomp.append(request.form.getlist(key))

    for key in request.form.keys():
        if key in column_names_pkcomp:
            column_data_pkcomp.append(request.form.getlist(key))

    for key in request.form.keys():
        if key in column_names_hours:
            column_data_hours.append(request.form.getlist(key))

    for key in request.form.keys():
        if key in column_names_comp:
            column_data_comp.append(request.form.getlist(key))

    for key in request.form.keys():
        if key in column_names_mto:
            column_data_mto.append(request.form.getlist(key))


    # Получение текущего количества строк в таблице

    num_rows_fullcomp = len(column_data_fullcomp[0])
    num_rows_ukcomp = len(column_data_ukcomp[0])
    num_rows_opkcomp = len(column_data_opkcomp[0])
    num_rows_pkcomp = len(column_data_pkcomp[0])
    num_rows_hours = len(column_data_hours[0])
    num_rows_comp = len(column_data_comp[0])
    num_rows_mto = len(column_data_mto[0])




    context = {
        'name_education': titul_name,
        'name_discipline': discipline,
        'school_name':school,
        'department': department,
        'name_direction': name_direction,
        'date_aprooval': date_aprooval,
        'form_education': form_education,
        'number_fgos': number_fgos,
        'name_author': authors,
        'year_education': year,
        'name_programm':program,
        'code_programm':code_education,
        'table_data_fullcomp': column_data_fullcomp,
        'table_data_uk':column_data_ukcomp,
        'table_data_opk':column_data_opkcomp,
        'table_data_pk':column_data_pkcomp,
        'table_hours':column_data_hours,
        'table_comp':column_data_comp,
        'table_mto':column_data_mto,
        'num_rows': num_rows_fullcomp,
        'form_control':form_control,
        'language_discipline':language_discipline,
        'target_discipline':target_discipline,
        'tasks_discipline':formatted_tasks,
        'list_competencies':list_competencies,
        'zachet_edenic':zachet_edenic,
        'akadem_hour':akadem_hour,
        'course':course,
        'theme_discipline':theme_discipline,
        'main_literature':main_literature,
        'additional_literature':additional_literature,
        'internet_literature':internet_literature,
        'it_literature':it_literature,
        'lecture_hour':lecture_hour,
        'practice_hour':practice_hour,
        'sm_hour':sm_hour
    }


    document.render(context)
    


    # # Получение таблицы из шаблона
    table_fullcomp = document.tables[1]
    table_ukcomp = document.tables[2]
    table_opkcomp = document.tables[3]
    table_pkcomp = document.tables[4]
    table_hours = document.tables[5]
    table_comp = document.tables[6]
    table_mto = document.tables[7]



    '''Все компетенции'''
    # Обновление таблицы с данными 
    current_rows_fullcomp = len(table_fullcomp.rows)
    if current_rows_fullcomp < num_rows_fullcomp:
        # Добавление недостающих строк
        for _ in range(num_rows_fullcomp - current_rows_fullcomp):
            table_fullcomp.add_row()

    # Обновление таблицы с данными
    for i in range(num_rows_fullcomp):
        row_cells = table_fullcomp.rows[i].cells
        for j, cell_data in enumerate(column_data_fullcomp):
            row_cells[j].text = cell_data[i]
            row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

   
    

    '''Универсальные компетенции'''
    current_rows_uk = len(table_ukcomp.rows)
    if current_rows_uk < num_rows_ukcomp:
        # Добавление недостающих строк
        for _ in range(num_rows_ukcomp - current_rows_uk):
            table_ukcomp.add_row()

    # Обновление таблицы с данными
    for i in range(num_rows_ukcomp):
        row_cells = table_ukcomp.rows[i].cells
        for j, cell_data in enumerate(column_data_ukcomp):
            row_cells[j].text = cell_data[i]
            row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    '''Общепрофессиональные компетенции'''
    current_rows_opk = len(table_opkcomp.rows)
    if current_rows_opk < num_rows_opkcomp:
        # Добавление недостающих строк
        for _ in range(num_rows_opkcomp - current_rows_opk):
            table_opkcomp.add_row()

    # Обновление таблицы с данными
    for i in range(num_rows_opkcomp):
        row_cells = table_opkcomp.rows[i].cells
        for j, cell_data in enumerate(column_data_opkcomp):
            row_cells[j].text = cell_data[i]
            row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    '''Профессиональные компетенции'''
    current_rows_pk = len(table_pkcomp.rows)
    if current_rows_pk < num_rows_pkcomp:
        # Добавление недостающих строк
        for _ in range(num_rows_pkcomp - current_rows_pk):
            table_pkcomp.add_row()

    # Обновление таблицы с данными
    for i in range(num_rows_pkcomp):
        row_cells = table_pkcomp.rows[i].cells
        for j, cell_data in enumerate(column_data_pkcomp):
            row_cells[j].text = cell_data[i]
            row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    '''Таблица с часами'''
    current_rows_hours = len(table_hours.rows)
    if current_rows_hours < num_rows_hours:
        # Добавление недостающих строк
        for i in range(num_rows_hours - current_rows_hours):
            table_hours.add_row()

    # Обновление таблицы с данными
    for i in range(num_rows_hours):
        row_cells_hour = table_hours.rows[i].cells
        for j, cell_data_hour in enumerate(column_data_hours):
            row_cells_hour[j].text = cell_data_hour[i]
            row_cells_hour[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    '''Таблица с темами и компетенциями'''
    current_rows_comp = len(table_comp.rows)
    if current_rows_comp < num_rows_comp:
        # Добавление недостающих строк
        for _ in range(num_rows_comp - current_rows_comp):
            table_comp.add_row()

    # Обновление таблицы с данными
    for i in range(num_rows_comp):
        row_cells = table_comp.rows[i].cells
        for j, cell_data in enumerate(column_data_comp):
            row_cells[j].text = cell_data[i]
            row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    '''Таблица с мто'''
    current_rows_mto = len(table_mto.rows)
    if current_rows_mto < num_rows_mto:
        # Добавление недостающих строк
        for i in range(num_rows_mto - current_rows_mto):
            table_mto.add_row()

    # Обновление таблицы с данными
    for i in range(num_rows_mto):
        row_cells = table_mto.rows[i].cells
        for j, cell_data in enumerate(column_data_mto):
            row_cells[j].text = cell_data[i]
            row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


    merge_cells_in_column(table_fullcomp, 1)
    merge_cells_in_column(table_ukcomp, 1)
    merge_cells_in_column(table_opkcomp, 1)
    merge_cells_in_column(table_pkcomp, 1)
    merge_cells_in_column(table_comp, 0)

    # Создание объекта для записи документа в память
    output = io.BytesIO()

    # Сохранение документа во временный файл
    document.save(output)
    output.seek(0)

    # Отправка содержимого файла клиенту
    return Response(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={'Content-Disposition': 'attachment;filename=document.docx'})


def merge_cells_in_column(table, column_index):
    rows = table.rows
    previous_value = None
    rowspan = 1

    for i in range(len(rows)):
        current_value = rows[i].cells[column_index].text

        if current_value == previous_value or current_value == "":
            rowspan += 1
            vMerge = rows[i].cells[column_index]._tc.get_or_add_tcPr().get_or_add_vMerge()
            vMerge.set(qn('w:val'), 'continue')

            if column_index > 0:
                vMerge_prev = rows[i].cells[column_index - 1]._tc.get_or_add_tcPr().get_or_add_vMerge()
                vMerge_prev.set(qn('w:val'), 'continue')
        else:
            if rowspan > 1:
                vMerge = rows[i - rowspan].cells[column_index]._tc.get_or_add_tcPr().get_or_add_vMerge()
                vMerge.set(qn('w:val'), 'restart')

                if column_index > 0:
                    vMerge_prev = rows[i - rowspan].cells[column_index - 1]._tc.get_or_add_tcPr().get_or_add_vMerge()
                    vMerge_prev.set(qn('w:val'), 'restart')
            previous_value = current_value
            rowspan = 1

        if i == len(rows) - 1 and rowspan > 1:
            vMerge = rows[i - rowspan + 1].cells[column_index]._tc.get_or_add_tcPr().get_or_add_vMerge()
            vMerge.set(qn('w:val'), 'restart')

            if column_index > 0:
                vMerge_prev = rows[i - rowspan + 1].cells[column_index - 1]._tc.get_or_add_tcPr().get_or_add_vMerge()
                vMerge_prev.set(qn('w:val'), 'restart')
